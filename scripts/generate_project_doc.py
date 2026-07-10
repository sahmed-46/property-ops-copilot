"""Generate Property Ops Copilot project Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Property_Ops_Copilot_Project_Documentation.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()

    title = doc.add_heading("Property Ops Copilot", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Multi-Agent PropTech Assistant with MCP Tool Servers")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Project Documentation | Syed Omar Ahmed | July 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "Property Ops Copilot is a multi-agent PropTech operations assistant that helps property "
        "managers answer lease questions, triage maintenance issues, and escalate compliance-sensitive "
        "cases. Specialist agents call custom MCP (Model Context Protocol) servers for lease retrieval, "
        "work order actions, and audit logging. The system is grounded (citations required), "
        "action-capable (ticket creation), and auditable (every tool call logged)."
    )
    add_para(doc, "One-line pitch: A property operations copilot that reads leases, creates tickets, and escalates safely—with a full audit trail.", bold=True)

    # 2. Problems
    add_heading(doc, "2. Problems Being Addressed", 1)
    add_table(doc, ["Problem", "Impact"], [
        ["Fragmented property operations", "PMs jump between lease PDFs, tickets, email, and spreadsheets"],
        ["Slow, inconsistent tenant/issue handling", "Urgent issues mis-prioritized; quality depends on who is on shift"],
        ["Lease knowledge trapped in documents", "Expensive, error-prone manual re-reading of leases"],
        ["Weak audit trail for AI-assisted decisions", "Leadership cannot verify what AI recommended or why"],
        ["Integration friction for agent systems", "Custom API glue per agent does not scale"],
    ])

    add_heading(doc, "2.1 Why Solving These Problems Is Required", 2)
    add_bullets(doc, [
        "Property operators face rising labor cost and portfolio complexity.",
        "Tenants expect faster resolution; delays hurt retention.",
        "Owners and investors face NOI and liability risk from SLA breaches and lease mistakes.",
        "PropTech needs connected decision workflows, not another generic chatbot.",
        "Compliance and legal teams require traceable handling for safety and lease obligations.",
    ])

    # 3. Benefits
    add_heading(doc, "3. Industry and Business Benefits", 1)
    add_heading(doc, "3.1 Industry Benefits", 2)
    add_bullets(doc, [
        "Practical pattern for agentic operations in real estate.",
        "MCP as a standard integration layer between AI agents and property systems.",
        "Human-in-the-loop adoption for high-risk cases (safety, legal, billing).",
    ])
    add_heading(doc, "3.2 Business Benefits", 2)
    add_table(doc, ["Metric", "How the Project Helps"], [
        ["Time to first response", "Auto-triage and draft replies"],
        ["Ticket routing accuracy", "Maintenance agent plus priority rules"],
        ["Lease lookup time", "RAG over clauses with citations"],
        ["Escalation miss rate", "Compliance agent flags risky cases"],
        ["Ops labor hours", "Fewer manual lookups and copy-paste"],
        ["Audit readiness", "MCP audit server logs every tool call"],
    ])

    # 4. Use Cases
    add_heading(doc, "4. Use Cases", 1)
    use_cases = [
        ("UC1: Lease Q&A with citations", "PM asks renewal date and notice requirement. Router → Lease Agent → MCP search_clauses + list_key_dates → cited answer."),
        ("UC2: Maintenance triage + work order", "Tenant reports no hot water. Router → Maintenance Agent → create_work_order → Comms drafts update."),
        ("UC3: Safety/compliance escalation", "Mold reported. Compliance Agent → check_policy → require_human_approval → block auto-close."),
        ("UC4: Ops briefing (stretch)", "Regional manager asks for open urgent tickets. Maintenance Agent → get_open_work_orders → SLA summary."),
        ("UC5: Agent observability", "Engineering lead inspects MCP tool calls, latency, failures, approval overrides."),
    ]
    for title, desc in use_cases:
        add_para(doc, title, bold=True)
        add_para(doc, desc)

    # 5. Is it one complete project
    add_heading(doc, "5. Is This One Complete Project?", 1)
    add_para(doc, "Yes. When scoped as Property Ops Copilot—a single product with one repo, one demo, and end-to-end flows—not as separate lease RAG, MCP, and maintenance demos.")
    add_para(doc, "MVP (portfolio-complete) includes:", bold=True)
    add_bullets(doc, [
        "3 MCP servers: lease-mcp, tickets-mcp, audit-mcp",
        "3–4 agents: Router, Lease, Maintenance, Compliance",
        "2 end-to-end flows: Lease Q&A and Maintenance ticket creation",
        "Audit/trace log of MCP tool calls",
        "GitHub README, architecture diagram, demo video",
    ])

    # 6. Scope
    add_heading(doc, "6. Project Scope", 1)
    add_heading(doc, "6.1 In Scope (MVP)", 2)
    add_bullets(doc, [
        "Chat UI for property manager",
        "Custom MCP servers for lease, tickets, and audit",
        "Sample data: 3 properties, 10 units, 10 tenants, 10 leases, 20 tickets",
        "Human approval gate for high-risk categories",
        "Eval set: 10–20 scripted scenarios",
    ])
    add_heading(doc, "6.2 Out of Scope (v1)", 2)
    add_bullets(doc, [
        "Real PMS integrations (Yardi, AppFolio, RealPage)",
        "Autonomous legal advice or auto-send without review",
        "Payments, evictions, accounting",
        "Mobile app and full enterprise RBAC",
    ])

    # 7. MCP
    add_heading(doc, "7. MCP Integration", 1)
    add_para(doc, "MCP serves as the standard tool layer between agents and property systems. Agents reason; MCP executes.")
    add_heading(doc, "7.1 Custom MCP Servers to Build", 2)
    add_table(doc, ["Server", "Tools"], [
        ["lease-mcp", "search_clauses, list_key_dates, get_lease_metadata, extract_lease_metadata"],
        ["tickets-mcp", "create_work_order, get_open_work_orders, update_work_order_status"],
        ["audit-mcp", "log_agent_action, require_human_approval, check_policy, get_trace"],
    ])
    add_heading(doc, "7.2 MCP Servers to Consume (optional)", 2)
    add_bullets(doc, [
        "Filesystem MCP for SOPs and templates",
        "PostgreSQL/SQLite MCP for property and ticket data",
        "Google Drive MCP for lease folders",
        "Slack MCP for PM notifications",
    ])

    add_heading(doc, "7.3 Multi-Agent Architecture", 2)
    add_para(doc, "User → Orchestrator (supervisor) → Router / Lease / Maintenance / Compliance / Comms agents → MCP Client Layer → lease-mcp, tickets-mcp, audit-mcp → PDFs, vector DB, SQLite, audit logs.")

    # 8. Tech Stack
    add_heading(doc, "8. Technology Stack", 1)
    add_table(doc, ["Layer", "Technology", "Purpose"], [
        ["Language", "Python 3.11+", "Core development"],
        ["API", "FastAPI", "REST backend"],
        ["UI", "Streamlit or React", "Chat and trace viewer"],
        ["Agents", "LangGraph or custom supervisor", "Multi-step orchestration"],
        ["LLM", "GPT-4o-mini / Claude Haiku", "Routing and synthesis"],
        ["Structured outputs", "Pydantic", "Tool args and agent decisions"],
        ["MCP", "Python MCP SDK", "Custom tool servers"],
        ["RAG", "Qdrant or FAISS + sentence-transformers", "Lease clause retrieval"],
        ["Database", "SQLite (MVP) → PostgreSQL", "Tickets, tenants, audit"],
        ["Documents", "pypdf + chunking", "Lease ingestion"],
        ["Testing", "pytest", "Tools, routing, escalation"],
        ["Deploy", "Docker Compose", "Reproducible demo"],
    ])

    # 9. Milestones
    add_heading(doc, "9. Milestones and Success Metrics", 1)
    add_table(doc, ["Week", "Deliverable"], [
        ["W1", "Sample data, lease ingestion, lease-mcp + tickets-mcp, unit tests"],
        ["W2", "Router + Lease + Maintenance agents, Lease Q&A end-to-end"],
        ["W3", "Compliance agent, audit-mcp, maintenance flow, escalation eval"],
        ["W4", "Trace UI, README, architecture diagram, demo video"],
    ])
    add_para(doc, "MVP Success Metrics:", bold=True)
    add_bullets(doc, [
        "≥80% lease Q&A accuracy on 15-scenario eval set",
        "≥90% maintenance routing accuracy on 10 scenarios",
        "100% recall on high-risk escalation test cases",
        "100% of MCP tool calls logged with timestamp and agent id",
    ])

    # 10. Build problems
    add_heading(doc, "10. Problems During Build", 1)
    add_bullets(doc, [
        "Messy tenant inputs (vague messages, wrong unit numbers)",
        "Ambiguous lease language (reasonable time, normal wear and tear)",
        "Overlapping intents (billing vs maintenance vs lease in one message)",
        "Poor PDF quality (scans, tables, inconsistent section numbering)",
        "Wrong agent or tool selection; invalid tool arguments",
        "Hallucinated citations; agent retry loops",
        "Hard debugging across multiple agents and MCP calls",
        "Non-deterministic routing; cost surprises from long contexts",
    ])

    # 11. Technical limitations
    add_heading(doc, "11. Technical Limitations", 1)
    add_table(doc, ["Area", "Limitation"], [
        ["LLM reasoning", "Probabilistic; not a licensed lawyer or PM"],
        ["RAG", "Misses clauses if chunking/embeddings are poor"],
        ["PDF parsing", "Tables and exhibits break extraction"],
        ["Latency", "Multi-agent + MCP can take 3–15+ seconds"],
        ["Context window", "Full lease + history can exceed limits"],
        ["MCP ecosystem", "Less mature than REST for auth and hosting"],
        ["SQLite MVP", "Not for concurrent production traffic"],
    ])

    # 12. Algorithm issues
    add_heading(doc, "12. Algorithm and AI Issues", 1)
    add_bullets(doc, [
        "Intent drift and class imbalance in routing",
        "RAG low recall/precision; chunk boundary errors",
        "Hallucinated citations; over-summarization of lease exceptions",
        "Wrong tool selection; malformed JSON tool arguments",
        "Error propagation and non-termination in multi-agent flows",
        "Vanity accuracy on easy FAQs while failing urgent edge cases",
    ])

    # 13. Guardrails
    add_heading(doc, "13. Guardrails Required", 1)
    layers = [
        ("Input", "Max length, file allowlist, PII redaction in logs, rate limits, prompt injection checks"),
        ("Routing", "Confidence threshold, agent allowlist, block action on ambiguous unit match"),
        ("Retrieval", "Citation required, span overlap check, active lease version only, score threshold"),
        ("Tool/MCP", "Pydantic validation, idempotency keys, role-based action allowlist"),
        ("Compliance rules", "Hard overrides for mold, no heat, legal threats; block unauthorized refund promises"),
        ("Output", "Structured format, forbidden phrases, human approval for outbound messages"),
        ("Audit", "Log agent, tools, args (redacted), results, latency, model version"),
        ("Human-in-the-loop", "Approve before send, close urgent tickets, legal summaries"),
    ]
    for name, desc in layers:
        add_para(doc, f"Layer: {name}", bold=True)
        add_para(doc, desc)

    # 14. Production traffic
    add_heading(doc, "14. Handling Production Traffic", 1)
    add_bullets(doc, [
        "Load balancer → API gateway (auth, rate limit) → request queue → worker pool",
        "Horizontally scale stateless MCP servers",
        "Cache frequent lease clauses; parallel MCP calls; async ticket creation",
        "Small model for routing; large model only for synthesis",
        "Queue burst traffic; return processing status via webhook/SSE",
        "Postgres read replicas; vector DB pre-filter by property",
        "Metrics: p50/p95 latency, tool error rate, cost per request",
    ])

    # 15. Secrets
    add_heading(doc, "15. Masking and Protecting Keys", 1)
    add_bullets(doc, [
        "Never commit API keys; use .env locally and Secret Manager in production",
        "MCP servers hold credentials; agents get scoped capability tokens",
        "Redact API keys, JWTs, SSN, bank info in logs (mask email/phone)",
        "Per-tool auth: tickets-mcp cannot delete leases",
        "Use gitleaks/trufflehog pre-commit; rotate keys if leaked",
        "Demo repos use synthetic data and spend-capped API keys only",
    ])

    # 16. Algorithms, skills, data
    add_heading(doc, "16. Algorithms, AI Techniques, Skills, and Data Formats", 1)
    add_heading(doc, "16.1 Algorithms and AI Techniques", 2)
    add_table(doc, ["Technique", "Use"], [
        ["LLMs", "Routing, synthesis, drafting"],
        ["RAG", "Lease clause Q&A"],
        ["Embeddings + vector search", "Semantic clause retrieval"],
        ["Intent classification", "Agent routing"],
        ["Structured output (Pydantic)", "Tool args and decisions"],
        ["Rule engine", "Escalation overrides"],
        ["Multi-agent orchestration", "Supervisor / LangGraph"],
        ["MCP tool calling", "Lease, ticket, audit actions"],
        ["Citation verification", "Grounding guardrail"],
    ])
    add_heading(doc, "16.2 Engineering and AI Skills", 2)
    add_para(doc, "Python, FastAPI, MCP, LangGraph, Pydantic, SQL, vector DBs, PDF parsing, Docker, pytest, prompt engineering, eval design, logging/observability, Git.")
    add_heading(doc, "16.3 Data Types and Formats", 2)
    add_table(doc, ["Data", "Format"], [
        ["Lease documents", "PDF → parsed Markdown/JSON clauses"],
        ["Properties, units, tenants, tickets", "SQL rows"],
        ["Embeddings", "Float vectors in Qdrant/FAISS"],
        ["Agent traces", "JSONL / SQL audit table"],
        ["MCP tool schemas", "JSON Schema"],
        ["Eval scenarios", "JSON/YAML"],
        ["Config", "YAML / .env"],
        ["API payloads", "JSON"],
    ])

    # 17. CV
    add_heading(doc, "17. CV and Portfolio Positioning", 1)
    add_para(doc, "Title: Property Ops Copilot — Multi-Agent PropTech Assistant (Python, FastAPI, MCP, RAG)", bold=True)
    add_para(doc, "This is a strong CV project when shipped with working MCP servers, 2 end-to-end flows, audit logs, and eval results. It demonstrates multi-agent orchestration, MCP creation, RAG grounding, guardrails, and production-minded engineering.")
    add_para(doc, "Sample resume bullet:", bold=True)
    add_para(
        doc,
        "Built a multi-agent property operations copilot with custom MCP servers for lease retrieval, "
        "work order management, and audit logging; achieved X% lease Q&A accuracy and 100% recall on "
        "high-risk escalation scenarios with human-in-the-loop approval gates."
    )
    add_para(doc, "Best combined with: AI Support Triage (structured outputs) and EcomAtlas (RAG depth) for a full applied AI portfolio.")

    # Definition of done
    add_heading(doc, "18. Definition of Done", 1)
    add_bullets(doc, [
        "GitHub repo with setup instructions",
        "3 MCP servers running locally",
        "4 agents with logged tool calls",
        "2 demo flows (lease Q&A + maintenance ticket)",
        "Eval doc with ≥15 scenarios and results",
        "Architecture diagram in README",
        "3–5 minute demo recording",
    ])

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
